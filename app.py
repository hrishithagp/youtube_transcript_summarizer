from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
from transformers import pipeline
from youtube_transcript_api import YouTubeTranscriptApi, NoTranscriptAvailable
from docx import Document
import os

app = Flask(__name__)
CORS(app)
model_name = "sshleifer/distilbart-cnn-12-6"
model_revision = "a4f8f3e"
summarizer = pipeline('summarization', model=model_name, revision=model_revision)

# Serve HTML page
@app.route('/')
def home():
    return render_template('index.html')

@app.route('/summarize', methods=['POST', 'OPTIONS'])
def summarize():
    if request.method == 'OPTIONS':
        response = jsonify({'status': 'ok'})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type')
        return response

    data = request.get_json()
    youtube_video = data['url']
    video_id = youtube_video.split("v=")[-1].split("&")[0]  # Handle additional URL parameters
    try:
        transcript = YouTubeTranscriptApi.get_transcript(video_id)
    except NoTranscriptAvailable:
        return jsonify({'error': 'Subtitles are not available for this video.'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 400

    result = []
    current_context = ""
    current_start_time = None

    for entry in transcript:
        text = entry['text']
        start_time = entry['start']
        end_time = entry['start'] + entry['duration']
        
        # Append text to current context
        if current_start_time is None:
            current_start_time = start_time
        current_context += text + " "

        # Check if the context length exceeds a certain limit (heuristic for context change)
        if len(current_context) > 1500:
            result.append({
                'context': current_context.strip(),
                'start_time': current_start_time,
                'end_time': end_time
            })
            current_context = ""
            current_start_time = None
    
    # Append the last context if there's any remaining
    if current_context:
        result.append({
            'context': current_context.strip(),
            'start_time': current_start_time,
            'end_time': end_time
        })

    # Summarize each context
    summarized_result = []
    for entry in result:
        summary = summarizer(entry['context'])[0]['summary_text']
        summarized_result.append({
            'summary': summary,
            'start_time': entry['start_time'],
            'end_time': entry['end_time']
        })

    # Save the summary to a Word document
    document = Document()
    document.add_heading('YouTube Video Summary', 0)
    for entry in summarized_result:
        document.add_paragraph(f"From {format_time(entry['start_time'])} to {format_time(entry['end_time'])}: {entry['summary']}")
    file_path = 'summary.docx'
    document.save(file_path)

    return jsonify({'summaries': summarized_result, 'doc_path': file_path})

@app.route('/download_summary', methods=['GET'])
def download_summary():
    file_path = request.args.get('doc_path')
    return send_file(file_path, as_attachment=True, download_name='summary.docx')

def format_time(seconds):
    minutes = int(seconds // 60)
    seconds = int(seconds % 60)
    return f"{minutes}:{seconds:02}"

if __name__ == "__main__":
    app.run(debug=True)
